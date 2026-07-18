"""POST /v1/policies/upload — policy document library with Gemma-driven compliance
gap analysis (tasks.md §17, Requirement 10.4).

Accepts a policy document (PDF/DOCX/TXT), extracts text, and asks Gemma to compare
it against the platform's own 40-rule compliance catalogue to surface gaps —
rules the policy doesn't address, or addresses incompletely. Falls back to a
keyword-coverage heuristic when Gemma is unavailable (never silently returns an
empty gap list without saying why).
"""
from __future__ import annotations

import json
import re
import uuid
from typing import Any, Dict, List, Optional

import structlog
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy import select, desc
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.gemma_client import get_llm_client_or_none, get_llm_model
from app.core.pii_redaction import redact_user_input
from app.core.security import get_db_user
from app.db.models import PolicyDocument
from app.db.session import get_db
from app.ml.rules_db import COMPLIANCE_RULES

log = structlog.get_logger()
router = APIRouter()

_ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class PolicyUploadResponse(BaseModel):
    policy_id: str
    title: str
    version: str
    compliance_gaps: List[Dict[str, Any]]
    gap_analysis_method: str  # "gemma" | "keyword_fallback"


def _file_ext(filename: str) -> str:
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


async def _extract_policy_text(file_bytes: bytes, filename: str) -> str:
    ext = _file_ext(filename)
    try:
        if ext == ".pdf":
            from app.ingestion.parse_pdf import extract_text_pages
            return "\n".join(extract_text_pages(file_bytes))
        elif ext == ".docx":
            from docx import Document as DocxDocument
            from io import BytesIO
            doc = DocxDocument(BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        else:
            return file_bytes.decode("utf-8", errors="replace")
    except Exception as exc:
        log.warning("policies.extract_failed", error=str(exc), filename=filename)
        return ""


GEMMA_GAP_PROMPT = """You are a compliance auditor. Below is a company policy document, followed by a
list of regulatory rules it should ideally address. Identify which rules the policy
does NOT adequately cover.

Return ONLY valid JSON: a list of objects like
[{{"rule_code": "AML-001", "gap_description": "Policy does not mention cash transaction reporting thresholds."}}]

Only include rules with a genuine gap — omit rules the policy already covers well.

RULES:
{rules_summary}

POLICY DOCUMENT:
{policy_text}
"""


def _keyword_fallback_gaps(policy_text: str) -> List[Dict[str, Any]]:
    """When Gemma is unavailable: flag any rule whose name/framework keywords
    don't appear anywhere in the policy text at all — a coarse but honest signal."""
    text_lower = policy_text.lower()
    gaps = []
    for rule in COMPLIANCE_RULES:
        keywords = [w.lower() for w in rule["name"].split() if len(w) > 4]
        if keywords and not any(kw in text_lower for kw in keywords):
            gaps.append({
                "rule_code": rule["code"],
                "gap_description": f"Policy text does not appear to reference '{rule['name']}' ({rule['framework']}) — verify manually.",
            })
    return gaps[:15]  # cap — this heuristic over-flags on short documents


@router.post("/policies/upload", response_model=PolicyUploadResponse)
async def upload_policy(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    version: str = Form("1.0"),
    effective_date: Optional[str] = Form(None),
    user=Depends(get_db_user),
    db: AsyncSession = Depends(get_db),
) -> PolicyUploadResponse:
    filename = file.filename or "policy"
    ext = _file_ext(filename)
    if ext not in _ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=422, detail=f"Unsupported file type '{ext}'. Accepted: pdf, docx, txt.")

    file_bytes = await file.read()
    raw_text = await _extract_policy_text(file_bytes, filename)
    if not raw_text.strip():
        raise HTTPException(status_code=422, detail="Could not extract any text from the uploaded file.")

    redaction = redact_user_input(raw_text, user["id"])
    safe_text = redaction.get("redacted_text", raw_text)

    client = get_llm_client_or_none()
    gaps: List[Dict[str, Any]] = []
    method = "keyword_fallback"

    if client is not None:
        try:
            rules_summary = "\n".join(f"- {r['code']}: {r['name']} ({r['framework']})" for r in COMPLIANCE_RULES)
            prompt = GEMMA_GAP_PROMPT.format(rules_summary=rules_summary, policy_text=safe_text[:6000])
            resp = client.chat.completions.create(
                model=get_llm_model(),
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=1024,
            )
            raw = resp.choices[0].message.content or "[]"
            raw = re.sub(r"^```(?:json)?\s*", "", raw.strip())
            raw = re.sub(r"\s*```$", "", raw.strip())
            gaps = json.loads(raw)
            method = "gemma"
        except Exception as exc:
            log.warning("policies.gemma_gap_analysis_failed", error=str(exc))
            gaps = _keyword_fallback_gaps(safe_text)
            method = "keyword_fallback"
    else:
        gaps = _keyword_fallback_gaps(safe_text)

    policy = PolicyDocument(
        user_id=uuid.UUID(user["db_id"]),
        title=title or filename,
        version=version,
        effective_date=effective_date,
        document_url=None,
        compliance_gaps=gaps,
    )
    db.add(policy)
    await db.commit()
    await db.refresh(policy)

    log.info("policies.upload_complete", policy_id=str(policy.id), gap_count=len(gaps), method=method)

    return PolicyUploadResponse(
        policy_id=str(policy.id),
        title=policy.title,
        version=policy.version or "1.0",
        compliance_gaps=gaps,
        gap_analysis_method=method,
    )


@router.get("/policies")
async def list_policies(
    limit: int = 50,
    user=Depends(get_db_user),
    db: AsyncSession = Depends(get_db),
) -> Dict[str, Any]:
    stmt = (
        select(PolicyDocument)
        .where(PolicyDocument.user_id == uuid.UUID(user["db_id"]))
        .order_by(desc(PolicyDocument.created_at))
        .limit(limit)
    )
    result = await db.execute(stmt)
    policies = result.scalars().all()

    return {
        "total": len(policies),
        "policies": [
            {
                "policy_id": str(p.id),
                "title": p.title,
                "version": p.version,
                "effective_date": p.effective_date,
                "compliance_gaps": p.compliance_gaps,
                "gap_count": len(p.compliance_gaps or []),
                "created_at": p.created_at.isoformat() if p.created_at else None,
            }
            for p in policies
        ],
    }
